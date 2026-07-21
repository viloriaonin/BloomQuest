import io
import re
import json
import logging
import uuid
import openpyxl
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from database import get_db
import models

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# Import AI utilities and our new layout parsing fallback
from ai_service import (
    generate_questions_from_tos,
    build_preview,
    prepare_database_rows,
    statistics,
    parse_syllabus_text_with_ai,
    _extract_ilo_label,
    GroqDailyQuotaExceeded,
)
from file_extractor import extract_text
from routers.tos_utils import (
    compute_tos,
    generate_tos_from_institutional_template,
)

router = APIRouter(prefix="/api/questions", tags=["Questions"])
logger = logging.getLogger(__name__)

FILE_CACHE = {}

# Pydantic schema for Table of Specifications payload
class TOSGenerationPayload(BaseModel):
    upload_id: str
    total_items: int
    whole_total_points: int
    question_types: list[str]
    selected_topic_indices: list[int]
    subcolumn_a_hours: dict[str, str]


def parse_syllabus_pdf(contents: bytes):
    """
    Layout-aware PDF text extractor using PyMuPDF4LLM, which handles
    table structure (including tables spanning multiple pages) far more
    reliably than manual page.get_text() stitching. Falls back to pypdf
    plain-text extraction if PyMuPDF4LLM isn't available or errors out.

    NOTE: we intentionally do NOT run a line-by-line regex parser over
    this text to detect Chapter/Unit + ILO pairs. Most CIS documents lay
    topics out as a table (Ch. | Topics | ILO columns), and a line regex
    has no way to tell which ILO belongs to which row once that table is
    flattened to text/markdown -- it ends up grabbing whichever
    "Outcome"/"ILO" line happens to appear nearest in the raw text, which
    is often the wrong one. pymupdf4llm already renders tables as
    structured markdown, so we hand that straight to the AI extractor,
    which is prompted to map each row's ILO to that same row's topic only.
    """
    full_text = ""
    try:
        import pymupdf4llm
        import fitz

        pdf = fitz.open(stream=contents, filetype="pdf")
        full_text = pymupdf4llm.to_markdown(pdf)

    except Exception as e:
        logger.warning(f"pymupdf4llm extraction failed, falling back to pypdf: {e}")
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(contents))
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text += text + "\n"

    logger.info(f"[SYLLABUS DEBUG] extracted full_text length={len(full_text)}")
    logger.info(f"[SYLLABUS DEBUG] extracted text sample:\n{full_text[:3000]}")

    return parse_syllabus_text_with_ai(full_text)


def _find_label_value(ws, label, max_row=200):
    label_lower = label.strip().lower()
    for row in ws.iter_rows(min_row=1, max_row=max_row):
        for cell in row:
            if cell.value is not None and str(cell.value).strip().lower() == label_lower:
                for c in range(cell.column + 1, cell.column + 15):
                    v = ws.cell(row=cell.row, column=c).value
                    if v is not None and str(v).strip():
                        return str(v).strip()
    return None

_SUBLINE_PREFIX_RE = re.compile(r"^[\-\u2022\*]\s*")
_READING_LIST_RE = re.compile(r"reading\s*list", re.IGNORECASE)


def _split_cell_lines(raw_val):
    """
    These CIS templates often pack the chapter title, its subtopics, and
    reading-list references into one wrapped Excel cell, separated by
    line breaks (subtopics usually prefixed with a dash/bullet).
    openpyxl returns that whole block as one string with embedded \n
    characters -- split it back into lines so we can tell the main
    topic apart from what's stacked underneath it.
    """
    if raw_val is None:
        return []
    return [line.strip() for line in str(raw_val).split("\n") if line.strip()]


def _main_topic_from_cell(raw_val):
    """Returns only the chapter/topic title -- the first line of the
    cell -- and drops every subtopic/tool/reading-list line stacked
    below it."""
    lines = _split_cell_lines(raw_val)
    if not lines:
        return ""
    return _SUBLINE_PREFIX_RE.sub("", lines[0]).strip()


def _ilo_from_cell(raw_val):
    """Returns only the actual outcome statement line(s) from the ILO
    cell, dropping any reading-list references bundled into the same
    wrapped cell alongside it."""
    lines = _split_cell_lines(raw_val)
    outcome_lines = [
        _SUBLINE_PREFIX_RE.sub("", ln).strip()
        for ln in lines
        if not _READING_LIST_RE.search(ln)
    ]
    return " ".join(outcome_lines).strip()

def _find_tla_header_row(ws, max_row=400):
    for row in ws.iter_rows(min_row=1, max_row=max_row):
        row_text = {cell.column: str(cell.value).strip() for cell in row if cell.value is not None}
        if not row_text:
            continue
        ch_col = next((c for c, v in row_text.items() if v.lower() in ("ch.", "ch")), None)
        topic_col = next((c for c, v in row_text.items() if "topics" in v.lower() and "reading" in v.lower()), None)
        if ch_col is not None and topic_col is not None:
            outcome_col = next((c for c, v in row_text.items() if "topic outcomes" in v.lower()), None)
            # Separate "ILO" header column -- distinct from "Topic
            # Outcomes". This CIS template maps each row to a short ILO
            # label (e.g. "ILO1") in its own column, apart from the
            # free-text outcome description. Match an exact "ILO" header
            # only, so it doesn't accidentally grab "Topic Outcomes" or
            # the "ILOs" mapping tables that appear later in the doc.
            ilo_col = next((c for c, v in row_text.items() if v.strip().lower() == "ilo"), None)
            return {
                "header_row": row[0].row,
                "ch_col": ch_col,
                "topic_col": topic_col,
                "outcome_col": outcome_col,
                "ilo_col": ilo_col,
            }
    return None

def _try_parse_number(val):
    if val is None: return None
    if isinstance(val, (int, float)): return float(val)
    try: return float(str(val).strip())
    except ValueError: return None

def parse_syllabus_excel(contents: bytes):
    workbook = openpyxl.load_workbook(io.BytesIO(contents), data_only=True)
    course_title, course_code, header_info, target_ws = None, None, None, None

    for sheet in workbook.worksheets:
        course_title = course_title or _find_label_value(sheet, "Course Title")
        course_code = course_code or _find_label_value(sheet, "Course Code")
        if header_info is None:
            found = _find_tla_header_row(sheet)
            if found:
                header_info = found
                target_ws = sheet

    if header_info is None or target_ws is None:
        full_text_lines = []
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip() != ""]
                if cells: full_text_lines.append(" ".join(cells))
        return parse_syllabus_text_with_ai("\n".join(full_text_lines))

    ws = target_ws
    ch_col = header_info["ch_col"]
    topic_col = header_info["topic_col"]
    outcome_col = header_info["outcome_col"]
    ilo_col = header_info.get("ilo_col")
    start_row = header_info["header_row"] + 1

    detected_topics = []
    for r in range(start_row, ws.max_row + 1):
        ch_val = ws.cell(row=r, column=ch_col).value
        topic_val = ws.cell(row=r, column=topic_col).value
        ch_num = _try_parse_number(ch_val)

        if ch_num is not None and topic_val and str(topic_val).strip():
            outcome_val = ws.cell(row=r, column=outcome_col).value if outcome_col else None
            ilo_val = ws.cell(row=r, column=ilo_col).value if ilo_col else None

            detected_topics.append({
                "name": _main_topic_from_cell(topic_val) or str(topic_val).strip(),
                "weight": 1.0,
                "ilo": _extract_ilo_label(str(ilo_val)) if ilo_val else "Not specified in CIS -- please review.",
                "ilo_description": _ilo_from_cell(outcome_val) if outcome_val else "",
            })
            continue
        if ch_val is not None and str(ch_val).strip() and ch_num is None:
            break

    return (course_title or "Dynamic Curricular Subject", course_code or "DYNAMIC-CODE", detected_topics)

# ============================================================
# ASSESSMENT DOCUMENT BUILDERS (docx / pdf)
# ============================================================

def _build_assessment_docx(questions, course_title, course_code) -> bytes:
    doc = Document()
    doc.add_heading(f"{course_code} - {course_title}", level=1)
    doc.add_paragraph("Examination")

    for i, q in enumerate(questions, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {q['question']}").bold = True
        if q.get("question_type") == "MCQ" and q.get("options"):
            for idx, opt in enumerate(q["options"]):
                doc.add_paragraph(f"    {chr(65 + idx)}. {opt}")

    doc.add_page_break()
    doc.add_heading("Answer Key", level=1)
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q['correct_answer']}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_assessment_pdf(questions, course_title, course_code) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER)
    styles = getSampleStyleSheet()
    story = [Paragraph(f"{course_code} - {course_title}", styles["Title"]), Spacer(1, 12)]

    for i, q in enumerate(questions, start=1):
        story.append(Paragraph(f"{i}. {q['question']}", styles["Normal"]))
        if q.get("question_type") == "MCQ" and q.get("options"):
            for idx, opt in enumerate(q["options"]):
                story.append(Paragraph(f"&nbsp;&nbsp;{chr(65 + idx)}. {opt}", styles["Normal"]))
        story.append(Spacer(1, 8))

    story.append(Spacer(1, 20))
    story.append(Paragraph("Answer Key", styles["Heading1"]))
    for i, q in enumerate(questions, start=1):
        story.append(Paragraph(f"{i}. {q['correct_answer']}", styles["Normal"]))

    doc.build(story)
    return buf.getvalue()


# --- STEP 1: UPLOAD ENDPOINT ---
@router.post("/upload")
async def upload_and_analyze_syllabus(
    module_file: UploadFile = File(...),
    syllabus_file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    upload_id = uuid.uuid4().hex
    filename = syllabus_file.filename.lower()

    try:
        module_contents = await module_file.read()
        module_text = extract_text(module_contents, module_file.filename)
        contents = await syllabus_file.read()

        if filename.endswith('.pdf'):
            course_title, course_code, detected_topics = parse_syllabus_pdf(contents)
        elif filename.endswith('.xlsx') or filename.endswith('.xls'):
            course_title, course_code, detected_topics = parse_syllabus_excel(contents)
        elif filename.endswith('.docx'):
            syllabus_text = extract_text(contents, syllabus_file.filename)
            course_title, course_code, detected_topics = parse_syllabus_text_with_ai(syllabus_text)
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported syllabus file format. Please upload a .pdf, .docx, .xlsx, or .xls file."
            )

        detected_subject = {
            "name": course_title,
            "code": course_code,
            "description": f"Dynamic dashboard tracker layout generated for {course_title}."
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Analysis Engine Error: {str(e)}")

    FILE_CACHE[f"{upload_id}_metadata"] = {"subject": detected_subject, "topics": detected_topics, "module_text": module_text}

    return {
        "upload_id": upload_id,
        "subject": detected_subject,
        "topics": detected_topics
    }


# --- STEP 2: MULTI-LEVEL TOS GENERATION AND DB CACHING ---
@router.post("/generate-with-tos")
async def generate_with_tos(
    payload: TOSGenerationPayload,
    db: Session = Depends(get_db),
):
    meta = FILE_CACHE.get(f"{payload.upload_id}_metadata")
    if not meta:
        raise HTTPException(status_code=404, detail="Upload session expired.")

    selected_topics_data = compute_tos(
        topics=meta["topics"],
        selected_topic_indices=payload.selected_topic_indices,
        hours_dict=payload.subcolumn_a_hours,
        total_items=payload.total_items,
        question_types=payload.question_types,
    )

    subject_row = db.query(models.Subject).filter(models.Subject.code == meta["subject"]["code"]).first()
    if not subject_row:
        subject_row = models.Subject(name=meta["subject"]["name"], code=meta["subject"]["code"])
        db.add(subject_row)
        db.commit()
        db.refresh(subject_row)

    try:
        generated_questions = generate_questions_from_tos(
            subject=meta["subject"],
            module_text=meta["module_text"],
            tos_data=selected_topics_data,
        )
    except GroqDailyQuotaExceeded as e:
        # Daily quota exhausted - not transient, so surface a clean
        # 429 with the actual wait time instead of a raw 502 dump.
        raise HTTPException(
            status_code=429,
            detail=str(e),
        )
    except Exception as e:
        logger.exception("Question generation failed while contacting the AI service")
        raise HTTPException(
            status_code=502,
            detail=f"Question generation failed while contacting the AI service. Details: {str(e)}",
        )

    rows = prepare_database_rows(generated_questions, subject_row.id)
    for row in rows:
        # ============================================================
        # POSTGRESQL ARRAY (TEXT[]) VALUE FORMATTING ENFORCEMENT
        # ============================================================
        correct_ans = row.get("correct_answer")
        if isinstance(correct_ans, str):
            correct_ans = [correct_ans]
        elif isinstance(correct_ans, dict):
            # Matching Type questions come back as {left_item: right_item, ...} --
            # psycopg2 can't adapt a raw dict into a TEXT[] column, so flatten it
            # into readable "left -> right" pair strings instead.
            correct_ans = [f"{left} -> {right}" for left, right in correct_ans.items()]
        elif isinstance(correct_ans, list):
            # Defensive: make sure every element is a plain string, in case the
            # AI nests dicts/numbers inside the list instead of a flat list of strings.
            flattened = []
            for item in correct_ans:
                if isinstance(item, dict):
                    flattened.extend(f"{k} -> {v}" for k, v in item.items())
                else:
                    flattened.append(str(item))
            correct_ans = flattened
        elif correct_ans is None:
            correct_ans = []
        else:
            # Catch-all for any other unexpected type (int, float, etc.)
            correct_ans = [str(correct_ans)]

        q = models.GeneratedQuestion(
            subject_id=row["subject_id"],
            topic_name=row["topic_name"],
            question=row["question"],
            bloom_level=row["bloom_level"],
            question_type=row["question_type"],
            options=row["options"],
            correct_answer=correct_ans,  # Swapped with our safely formatted array structure
            explanation=row["explanation"],
        )
        db.add(q)
    db.commit()

    FILE_CACHE[f"{payload.upload_id}_questions"] = generated_questions

    workbook = generate_tos_from_institutional_template(
        selected_topics_data=selected_topics_data,
        course_code=subject_row.code,
        course_title=subject_row.name,
        whole_total_points=payload.whole_total_points,
    )
    stream = io.BytesIO()
    workbook.save(stream)
    FILE_CACHE[f"{payload.upload_id}_tos"] = stream.getvalue()

    preview = build_preview(generated_questions)
    stats = statistics(generated_questions)

    def _build_tos_response_rows(tos_data):
        tos_rows = []
        for topic in tos_data:
            bloom_breakdown = {
                bloom: {"total": count}
                for bloom, count in topic.get("bloom_counts", {}).items()
            }
            tos_rows.append({
                "topic": topic.get("topic_name", ""),
                "weight": topic.get("weight", 0),
                "total_items": topic.get("items", 0),
                "bloom_breakdown": bloom_breakdown,
                "question_distribution": topic.get("question_distribution", {}),
                "ilo": topic.get("ilo", ""),
                "ilo_description": topic.get("ilo_description", ""),
            })
        return tos_rows

    return {
        "message": "Assessment generated successfully.",
        "tos": _build_tos_response_rows(selected_topics_data),
        "questions_preview": preview,
        "statistics": stats,
        "total_questions": len(generated_questions),
    }


# --- STEP 3: EXPORT ROUTES ---

@router.get("/export/tos")
async def export_institutional_tos(upload_id: str):
    """
    Retrieves the generated openpyxl Excel spreadsheet payload matching
    the active session token directly from the shared memory cache.
    """
    tos_binary = FILE_CACHE.get(f"{upload_id}_tos")

    if not tos_binary:
        raise HTTPException(
            status_code=404,
            detail="TOS file asset records not found or the cache session has expired."
        )

    return StreamingResponse(
        io.BytesIO(tos_binary),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f"attachment; filename=Institutional_TOS_{upload_id[:8]}.xlsx",
            "Access-Control-Expose-Headers": "Content-Disposition"
        }
    )


@router.get("/export/assessment/docx")
async def export_assessment_docx(upload_id: str):
    questions = FILE_CACHE.get(f"{upload_id}_questions")
    meta = FILE_CACHE.get(f"{upload_id}_metadata")
    if not questions or not meta:
        raise HTTPException(status_code=404, detail="Generated assessment not found or session expired.")

    data = _build_assessment_docx(questions, meta["subject"]["name"], meta["subject"]["code"])
    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=Exam_Paper_With_Keys.docx",
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )


@router.get("/export/assessment/pdf")
async def export_assessment_pdf(upload_id: str):
    questions = FILE_CACHE.get(f"{upload_id}_questions")
    meta = FILE_CACHE.get(f"{upload_id}_metadata")
    if not questions or not meta:
        raise HTTPException(status_code=404, detail="Generated assessment not found or session expired.")

    data = _build_assessment_pdf(questions, meta["subject"]["name"], meta["subject"]["code"])
    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/pdf",
        headers={
            "Content-Disposition": "attachment; filename=Exam_Paper_With_Keys.pdf",
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )