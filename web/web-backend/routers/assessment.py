# routers/assessment.py
from fastapi import APIRouter, HTTPException, Depends, Form
from fastapi.responses import FileResponse
from fastapi.background import BackgroundTasks
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, uuid, tempfile
import pythoncom
from docx2pdf import convert
from database import get_db
import models

router = APIRouter(prefix="/api/assessment", tags=["Assessment"])

# Second router so the path matches exactly what the frontend calls:
# POST /api/questions/export
export_router = APIRouter(prefix="/api/questions", tags=["Questions"])

TEMP_DIR = tempfile.gettempdir()


def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    pythoncom.CoInitialize()
    try:
        convert(docx_path, pdf_path)
    finally:
        pythoncom.CoUninitialize()


def build_assessment_docx(subject: models.Subject, questions: list) -> str:
    doc = Document()

    # ── Title ──
    title = doc.add_heading(f"{subject.name} — Assessment", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Total Items: {len(questions)}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Name: ____________________    Score: _______")
    doc.add_paragraph()

    # ── Questions ──
    letters = ['A', 'B', 'C', 'D', 'E', 'F']
    for i, q in enumerate(questions, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {q.question}").bold = True

        if q.options:
            for j, opt in enumerate(q.options):
                doc.add_paragraph(f"   {letters[j]}. {opt}")
        else:
            doc.add_paragraph("   Answer: ____________________________________")
        doc.add_paragraph()

    # ── Answer Key (new page) ──
    doc.add_page_break()
    doc.add_heading("Answer Key", level=1)

    for i, q in enumerate(questions, start=1):
        answer = q.correct_answer or "N/A"
        doc.add_paragraph(f"{i}. {answer}")

    file_path = os.path.join(TEMP_DIR, f"assessment_{uuid.uuid4().hex}.docx")
    doc.save(file_path)
    return file_path


def cleanup_file(path: str):
    if os.path.exists(path):
        os.remove(path)


def _get_questions_by_ids(db: Session, subject_id: int, question_ids: list[int]):
    subject = db.query(models.Subject).filter(models.Subject.id == subject_id).first()
    if not subject:
        raise HTTPException(404, "Subject not found")

    questions = (
        db.query(models.GeneratedQuestion)
        .filter(
            models.GeneratedQuestion.subject_id == subject_id,
            models.GeneratedQuestion.id.in_(question_ids),
        )
        .all()
    )
    if not questions:
        raise HTTPException(404, "No matching questions found for this subject")

    return subject, questions


# ──────────────────────────────────────────────────────────────
# Original endpoint (GET, generates from the WHOLE subject).
# Left in place in case anything else relies on it.
# ──────────────────────────────────────────────────────────────
@router.get("/generate")
def generate_assessment(
    subject_id: int,
    file_type: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
):
    if file_type not in ("docx", "pdf"):
        raise HTTPException(400, "file_type must be 'docx' or 'pdf'")

    subject = db.query(models.Subject).filter(models.Subject.id == subject_id).first()
    if not subject:
        raise HTTPException(404, "Subject not found")

    questions = (
        db.query(models.GeneratedQuestion)
        .filter(models.GeneratedQuestion.subject_id == subject_id)
        .all()
    )
    if not questions:
        raise HTTPException(404, "No questions found for this subject")

    docx_path = build_assessment_docx(subject, questions)
    safe_name = subject.name.replace(' ', '_')

    if file_type == "docx":
        background_tasks.add_task(cleanup_file, docx_path)
        return FileResponse(
            docx_path,
            filename=f"{safe_name}_Assessment.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # pdf
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert_docx_to_pdf(docx_path, pdf_path)

    background_tasks.add_task(cleanup_file, docx_path)
    background_tasks.add_task(cleanup_file, pdf_path)
    return FileResponse(
        pdf_path,
        filename=f"{safe_name}_Assessment.pdf",
        media_type="application/pdf",
    )


# ──────────────────────────────────────────────────────────────
# New endpoint matching questionbank.jsx exactly:
# POST /api/questions/export
# FormData: subject_id, question_ids ("1,4,7"), export_format ("pdf" | "docx")
# ──────────────────────────────────────────────────────────────
@export_router.post("/export")
def export_selected_questions(
    background_tasks: BackgroundTasks,
    subject_id: int = Form(...),
    question_ids: str = Form(...),
    export_format: str = Form("pdf"),
    db: Session = Depends(get_db),
):
    if export_format not in ("docx", "pdf"):
        raise HTTPException(400, "export_format must be 'docx' or 'pdf'")

    try:
        id_list = [int(qid) for qid in question_ids.split(",") if qid.strip() != ""]
    except ValueError:
        raise HTTPException(400, "question_ids must be a comma-separated list of integers")

    if not id_list:
        raise HTTPException(400, "No question_ids provided")

    subject, questions = _get_questions_by_ids(db, subject_id, id_list)

    docx_path = build_assessment_docx(subject, questions)
    safe_name = subject.name.replace(' ', '_')

    if export_format == "docx":
        background_tasks.add_task(cleanup_file, docx_path)
        return FileResponse(
            docx_path,
            filename=f"{safe_name}_Assessment.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # pdf
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert_docx_to_pdf(docx_path, pdf_path)

    background_tasks.add_task(cleanup_file, docx_path)
    background_tasks.add_task(cleanup_file, pdf_path)
    return FileResponse(
        pdf_path,
        filename=f"{safe_name}_Assessment.pdf",
        media_type="application/pdf",
    )